from docx import Document
from docx.shared import Pt, Inches, RGBColor

doc = Document('DOC TO REFER TO/ASSESSMENT PLAN Template .docx')

print("=== DOCUMENT ANALYSIS ===\n")

# Page setup
section = doc.sections[0]
print(f"Page width: {section.page_width.inches:.2f} inches")
print(f"Page height: {section.page_height.inches:.2f} inches")
print(f"Orientation: {'Landscape' if section.page_width > section.page_height else 'Portrait'}")
print(f"Top margin: {section.top_margin.inches:.2f} inches")
print(f"Bottom margin: {section.bottom_margin.inches:.2f} inches")
print(f"Left margin: {section.left_margin.inches:.2f} inches")
print(f"Right margin: {section.right_margin.inches:.2f} inches\n")

# Paragraphs
print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"\nPara {i}: '{para.text}'")
        if para.runs:
            run = para.runs[0]
            print(f"  Font: {run.font.name}")
            print(f"  Size: {run.font.size.pt if run.font.size else 'None'} pt")
            print(f"  Bold: {run.font.bold}")
            print(f"  Alignment: {para.alignment}")

# Table
print("\n\n=== TABLE ===")
table = doc.tables[0]
print(f"Rows: {len(table.rows)}")
print(f"Columns: {len(table.rows[0].cells)}")

print("\n--- Header Row ---")
for i, cell in enumerate(table.rows[0].cells):
    print(f"Col {i}: '{cell.text.strip()}'")
    if cell.paragraphs and cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        print(f"  Font: {run.font.name}, Size: {run.font.size.pt if run.font.size else 'None'} pt, Bold: {run.font.bold}")
    # Check cell shading
    try:
        shading = cell._element.xpath('.//w:shd')[0]
        fill = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
        print(f"  Background: {fill}")
    except:
        pass

print("\n--- Column Widths ---")
for i, cell in enumerate(table.rows[0].cells):
    print(f"Col {i} width: {cell.width.inches:.2f} inches")
