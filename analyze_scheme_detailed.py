from docx import Document

doc = Document('DOC TO REFER TO/#2_schemeofwork_C++.docx')

print("=" * 80)
print("RTB SCHEME OF WORK - OFFICIAL TEMPLATE ANALYSIS")
print("=" * 80)

# Analyze header info table
print("\nHEADER TABLE (Table 0):")
table0 = doc.tables[0]
for i, row in enumerate(table0.rows):
    print(f"Row {i}: {' | '.join([cell.text.strip()[:50] for cell in row.cells])}")

# Analyze main content table
print("\nMAIN CONTENT TABLE (Table 1):")
table1 = doc.tables[1]
print(f"Dimensions: {len(table1.rows)} rows x {len(table1.columns)} columns")
print("\nHeader Row 0:")
for i, cell in enumerate(table1.rows[0].cells):
    print(f"  Col {i}: {cell.text.strip()}")
print("\nHeader Row 1:")
for i, cell in enumerate(table1.rows[1].cells):
    print(f"  Col {i}: {cell.text.strip()}")

print("\nDATA ROWS:")
for i in range(2, min(5, len(table1.rows))):
    print(f"\nRow {i}:")
    for j, cell in enumerate(table1.rows[i].cells):
        text = cell.text.strip()[:60]
        if text:
            print(f"  Col {j}: {text}")

print("\n" + "=" * 80)
