from docx import Document

doc = Document('DOC TO REFER TO/#1_session_plan1.docx')
print('=== FACILITATION TECHNIQUE ===')
print(doc.tables[0].rows[8].cells[0].text)
print('\n=== DEVELOPMENT SECTION ===')
for i in range(11, 16):
    print(f'Row {i}:', doc.tables[0].rows[i].cells[1].text[:300])
