"""
RTB Document Generator - Official Format
Matches official RTB session plan and scheme templates exactly
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import tempfile

# Font: Times New Roman 12pt (official RTB standard)
RTB_FONT = 'Times New Roman'
RTB_FONT_SIZE = 12

def set_cell_text(cell, text, bold=False, font_size=RTB_FONT_SIZE):
    """Set cell text with Times New Roman"""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = RTB_FONT
    run.font.bold = bold
    return cell

def generate_apa_references(topic, sector, learning_outcomes):
    """Generate real APA format references"""
    topic_lower = topic.lower()
    sector_lower = sector.lower() if sector else ''
    
    refs = []
    
    if 'ict' in sector_lower or 'software' in sector_lower or 'programming' in topic_lower:
        refs.extend([
            "Deitel, P., & Deitel, H. (2020). C++ how to program (10th ed.). Pearson.",
            "Sommerville, I. (2021). Software engineering (10th ed.). Pearson Education.",
            "Silberschatz, A., Galvin, P. B., & Gagne, G. (2018). Operating system concepts (10th ed.). Wiley."
        ])
    elif 'hospitality' in sector_lower or 'food' in sector_lower:
        refs.extend([
            "Gisslen, W. (2018). Professional cooking (9th ed.). Wiley.",
            "Walker, J. R., & Walker, J. T. (2020). Introduction to hospitality management (5th ed.). Pearson."
        ])
    else:
        refs.extend([
            "Rwanda TVET Board. (2022). Competency-based training curriculum guidelines. RTB.",
            "UNESCO-UNEVOC. (2021). TVET teaching and learning: A practical guide. UNESCO."
        ])
    
    return "\\n".join(refs[:4])

def generate_smart_objectives(topic, learning_outcomes):
    """Generate SMART objectives"""
    lo_lower = learning_outcomes.lower()
    
    if any(word in lo_lower for word in ['analyze', 'evaluate', 'create']):
        verbs = ['Analyze', 'Evaluate', 'Create']
    elif any(word in lo_lower for word in ['apply', 'demonstrate', 'implement']):
        verbs = ['Demonstrate', 'Apply', 'Implement']
    else:
        verbs = ['Define', 'Identify', 'Explain']
    
    objectives = [
        f"{verbs[0]} {topic} and its key concepts.",
        f"{verbs[1]} practical skills related to {topic}.",
        f"{verbs[2]} knowledge to solve problems in {topic}."
    ]
    
    return "\\n".join([f"{i+1}. {obj}" for i, obj in enumerate(objectives)])

def generate_session_plan_docx(session_plan):
    """Generate RTB session plan - OFFICIAL FORMAT"""
    doc = Document()
    
    # LANDSCAPE orientation (official RTB)
    section = doc.sections[0]
    section.orientation = 1
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('SESSION PLAN')
    title_run.font.name = RTB_FONT
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    # Main table: 22 rows x 8 columns (OFFICIAL RTB)
    table = doc.add_table(rows=22, cols=8)
    table.style = 'Table Grid'
    
    # Helper function
    def set_text(row, col, text, bold=True, colspan=1):
        cell = table.cell(row, col)
        if colspan > 1:
            cell.merge(table.cell(row, col + colspan - 1))
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.name = RTB_FONT
        run.font.size = Pt(12)
        run.font.bold = bold
        return cell
    
    # Row 0: Sector, Trade, Date
    set_text(0, 0, f"Sector: {session_plan.sector or ''}", colspan=2)
    set_text(0, 2, f"Trade: {session_plan.trade or ''}", colspan=4)
    set_text(0, 6, f"Date: {session_plan.date or ''}", colspan=2)
    
    # Row 1: Trainer, Term
    set_text(1, 0, f"Lead Trainer's name: {session_plan.trainer_name or ''}", colspan=5)
    set_text(1, 5, f"TERM: {session_plan.term or ''}", colspan=3)
    
    # Row 2: Module, Week, Trainees, Class
    set_text(2, 0, f"Module: {session_plan.module_code_title or ''}", colspan=3)
    set_text(2, 3, f"Week: {session_plan.week or ''}", colspan=2)
    set_text(2, 5, f"No. Trainees: {session_plan.number_of_trainees or ''}")
    set_text(2, 6, f"Class: {session_plan.class_name or ''}", colspan=2)
    
    # Row 3: Learning outcome
    set_text(3, 0, "Learning outcome:", colspan=2)
    set_text(3, 2, session_plan.learning_outcomes or '', bold=False, colspan=6)
    
    # Row 4: Indicative contents
    set_text(4, 0, "Indicative contents:", colspan=2)
    set_text(4, 2, session_plan.indicative_contents or '', bold=False, colspan=6)
    
    # Row 5: Topic
    set_text(5, 0, f"Topic: {session_plan.topic_of_session or ''}", colspan=8)
    
    # Row 6: Duration
    set_text(6, 0, f"Duration: {session_plan.duration or '40'} minutes", colspan=8)
    
    # Row 7: Objectives
    objectives = generate_smart_objectives(
        session_plan.topic_of_session or 'the topic',
        session_plan.learning_outcomes or ''
    )
    set_text(7, 0, "Objectives:", colspan=2)
    set_text(7, 2, objectives, bold=False, colspan=6)
    
    # Row 8: Facilitation technique
    set_text(8, 0, f"Facilitation technique: {session_plan.facilitation_techniques or 'Trainer Guided'}", colspan=8)
    
    # Row 9: Headers
    set_text(9, 0, "Introduction")
    set_text(9, 1, "Activities", colspan=5)
    set_text(9, 6, "Resources")
    set_text(9, 7, "Duration")
    
    # Row 10-14: Development (technique-based)
    technique = session_plan.facilitation_techniques or 'Trainer Guided'
    topic = session_plan.topic_of_session or 'the topic'
    
    if technique == 'Brainstorming':
        activities = f"Brainstorming on {topic}:\\n- Generate ideas\\n- Cluster concepts\\n- Apply to practice"
    elif technique == 'Trainer Guided':
        activities = f"Trainer Guided on {topic}:\\n- I Do: Demonstration\\n- We Do: Guided practice\\n- You Do: Independent work"
    elif technique == 'Group Discussion':
        activities = f"Group Discussion on {topic}:\\n- Present question\\n- Small group work\\n- Whole class sharing"
    else:
        activities = f"Learning activities on {topic}:\\n- Introduction\\n- Practice\\n- Application"
    
    set_text(10, 0, "Development")
    set_text(10, 1, activities, bold=False, colspan=5)
    set_text(10, 6, "Computers\\nProjector\\nMaterials", bold=False)
    set_text(10, 7, "25 min")
    
    # Merge rows 11-14 for development
    for row in range(11, 15):
        table.cell(row, 0).merge(table.cell(10, 0))
        table.cell(row, 1).merge(table.cell(10, 1))
        table.cell(row, 6).merge(table.cell(10, 6))
        table.cell(row, 7).merge(table.cell(10, 7))
    
    # Row 15: Conclusion
    set_text(15, 0, "Conclusion")
    set_text(15, 1, "Summary and review", bold=False, colspan=5)
    set_text(15, 6, "Projector", bold=False)
    set_text(15, 7, "5 min")
    
    # Row 16-18: Assessment
    set_text(16, 0, "Assessment")
    set_text(16, 1, f"Practical tasks on {topic}", bold=False, colspan=5)
    set_text(16, 6, "Assessment sheets", bold=False)
    set_text(16, 7, "5 min")
    
    for row in range(17, 19):
        table.cell(row, 0).merge(table.cell(16, 0))
        table.cell(row, 1).merge(table.cell(16, 1))
        table.cell(row, 6).merge(table.cell(16, 6))
        table.cell(row, 7).merge(table.cell(16, 7))
    
    # Row 19: References
    refs = generate_apa_references(topic, session_plan.sector or '', session_plan.learning_outcomes or '')
    set_text(19, 0, "References:", colspan=2)
    set_text(19, 2, refs, bold=False, colspan=6)
    
    # Row 20: Appendices
    set_text(20, 0, "Appendices:", colspan=2)
    set_text(20, 2, f"PowerPoint on {topic}, Worksheets, Assessment rubrics", bold=False, colspan=6)
    
    # Row 21: Reflection
    set_text(21, 0, "Reflection:", colspan=2)
    set_text(21, 2, session_plan.reflection or f"Session on {topic} delivered successfully.", bold=False, colspan=6)
    
    # Save
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

# Keep existing scheme generator
def generate_scheme_of_work_docx(scheme):
    """Generate scheme of work"""
    # Import from original file
    from document_generator import generate_scheme_of_work_docx as original_scheme
    return original_scheme(scheme)
