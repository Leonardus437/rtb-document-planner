"""
RTB Document Generator - BRILLIANT OFFICIAL FORMAT
Creates absolutely amazing RTB-compliant documents that will impress TVET teachers
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import tempfile

# Font: Bookman Old Style 12pt (official RTB standard)
RTB_FONT = 'Bookman Old Style'
RTB_FONT_SIZE = 12

def set_cell_text(cell, text, bold=False, font_size=RTB_FONT_SIZE):
    """Set cell text with Bookman Old Style and handle line breaks"""
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = 1.5  # Set 1.5 line spacing
    lines = str(text).split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            para.add_run('\n')
        run = para.add_run(line)
        run.font.size = Pt(font_size)
        run.font.name = RTB_FONT
        run.font.bold = bold
    return cell

def set_cell_color(cell, color_hex):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)

def generate_apa_references(topic, sector, learning_outcomes):
    """Generate dynamic APA-format references based on specific topic"""
    topic_lower = (topic or '').lower()
    sector_lower = (sector or '').lower()
    topic_clean = (topic or 'the topic').title()
    
    refs = []
    
    # Programming & Software topics
    if any(word in topic_lower for word in ['python', 'java', 'c++', 'programming', 'coding', 'software']):
        if 'python' in topic_lower:
            refs = [
                f"Lutz, M. (2021). Learning Python: Powerful object-oriented programming (5th ed.). O'Reilly Media.",
                f"McKinney, W. (2022). Python for data analysis (3rd ed.). O'Reilly Media.",
                f"Matthes, E. (2019). Python crash course: A hands-on, project-based introduction to programming (2nd ed.). No Starch Press.",
                f"Rwanda TVET Board. (2023). {topic_clean} training curriculum. RTB."
            ]
        elif 'java' in topic_lower:
            refs = [
                f"Horstmann, C. S. (2019). Core Java Volume I: Fundamentals (11th ed.). Prentice Hall.",
                f"Sierra, K., & Bates, B. (2020). Head First Java (3rd ed.). O'Reilly Media.",
                f"Bloch, J. (2018). Effective Java (3rd ed.). Addison-Wesley.",
                f"Rwanda TVET Board. (2023). {topic_clean} competency standards. RTB."
            ]
        elif 'c++' in topic_lower or 'cpp' in topic_lower:
            refs = [
                f"Stroustrup, B. (2020). The C++ programming language (4th ed.). Addison-Wesley.",
                f"Deitel, P., & Deitel, H. (2020). C++ how to program (10th ed.). Pearson.",
                f"Lippman, S. B., Lajoie, J., & Moo, B. E. (2019). C++ primer (5th ed.). Addison-Wesley.",
                f"Rwanda TVET Board. (2023). {topic_clean} training manual. RTB."
            ]
        else:
            refs = [
                f"Sommerville, I. (2021). Software engineering (10th ed.). Pearson Education.",
                f"Pressman, R. S., & Maxim, B. R. (2020). Software engineering: A practitioner's approach (9th ed.). McGraw-Hill.",
                f"Martin, R. C. (2019). Clean code: A handbook of agile software craftsmanship. Prentice Hall.",
                f"Rwanda TVET Board. (2023). {topic_clean} curriculum guidelines. RTB."
            ]
    
    # Database topics
    elif any(word in topic_lower for word in ['database', 'sql', 'mysql', 'oracle', 'data']):
        refs = [
            f"Elmasri, R., & Navathe, S. B. (2021). Fundamentals of database systems (7th ed.). Pearson.",
            f"Silberschatz, A., Korth, H. F., & Sudarshan, S. (2020). Database system concepts (7th ed.). McGraw-Hill.",
            f"Date, C. J. (2019). An introduction to database systems (8th ed.). Addison-Wesley.",
            f"Rwanda TVET Board. (2023). {topic_clean} training standards. RTB."
        ]
    
    # Networking topics
    elif any(word in topic_lower for word in ['network', 'cisco', 'routing', 'switching', 'internet']):
        refs = [
            f"Kurose, J. F., & Ross, K. W. (2021). Computer networking: A top-down approach (8th ed.). Pearson.",
            f"Tanenbaum, A. S., & Wetherall, D. J. (2020). Computer networks (6th ed.). Pearson.",
            f"Odom, W. (2019). CCNA 200-301 official cert guide library. Cisco Press.",
            f"Rwanda TVET Board. (2023). {topic_clean} competency framework. RTB."
        ]
    
    # Web Development topics
    elif any(word in topic_lower for word in ['web', 'html', 'css', 'javascript', 'react', 'angular']):
        refs = [
            f"Duckett, J. (2021). HTML and CSS: Design and build websites. Wiley.",
            f"Flanagan, D. (2020). JavaScript: The definitive guide (7th ed.). O'Reilly Media.",
            f"Robbins, J. N. (2019). Learning web design: A beginner's guide (5th ed.). O'Reilly Media.",
            f"Rwanda TVET Board. (2023). {topic_clean} training curriculum. RTB."
        ]
    
    # Hospitality & Tourism
    elif any(word in topic_lower for word in ['hospitality', 'hotel', 'tourism', 'food', 'culinary', 'cooking', 'restaurant']):
        refs = [
            f"Gisslen, W. (2018). Professional cooking (9th ed.). Wiley.",
            f"Walker, J. R. (2020). Introduction to hospitality management (5th ed.). Pearson.",
            f"Barrows, C. W., & Powers, T. (2019). Introduction to management in the hospitality industry (11th ed.). Wiley.",
            f"Rwanda TVET Board. (2023). {topic_clean} training standards. RTB."
        ]
    
    # Construction & Building
    elif any(word in topic_lower for word in ['construction', 'building', 'architecture', 'civil', 'concrete', 'masonry']):
        refs = [
            f"Ching, F. D. K. (2020). Building construction illustrated (6th ed.). Wiley.",
            f"Allen, E., & Iano, J. (2019). Fundamentals of building construction (7th ed.). Wiley.",
            f"Neville, A. M. (2018). Properties of concrete (5th ed.). Pearson Education.",
            f"Rwanda TVET Board. (2023). {topic_clean} curriculum guidelines. RTB."
        ]
    
    # Electrical & Electronics
    elif any(word in topic_lower for word in ['electrical', 'electronics', 'circuit', 'wiring', 'power']):
        refs = [
            f"Boylestad, R. L. (2020). Introductory circuit analysis (13th ed.). Pearson.",
            f"Floyd, T. L. (2019). Principles of electric circuits (10th ed.). Pearson.",
            f"Malvino, A. P., & Bates, D. J. (2021). Electronic principles (9th ed.). McGraw-Hill.",
            f"Rwanda TVET Board. (2023). {topic_clean} training manual. RTB."
        ]
    
    # Mechanical & Automotive
    elif any(word in topic_lower for word in ['mechanical', 'automotive', 'engine', 'vehicle', 'machine']):
        refs = [
            f"Heisler, H. (2019). Advanced vehicle technology (3rd ed.). Butterworth-Heinemann.",
            f"Shigley, J. E., & Mischke, C. R. (2020). Mechanical engineering design (11th ed.). McGraw-Hill.",
            f"Erjavec, J. (2021). Automotive technology: A systems approach (7th ed.). Cengage Learning.",
            f"Rwanda TVET Board. (2023). {topic_clean} competency standards. RTB."
        ]
    
    # Business & Management
    elif any(word in topic_lower for word in ['business', 'management', 'entrepreneurship', 'marketing', 'accounting']):
        refs = [
            f"Robbins, S. P., & Coulter, M. (2021). Management (15th ed.). Pearson.",
            f"Kotler, P., & Armstrong, G. (2020). Principles of marketing (18th ed.). Pearson.",
            f"Daft, R. L. (2019). Management (13th ed.). Cengage Learning.",
            f"Rwanda TVET Board. (2023). {topic_clean} training curriculum. RTB."
        ]
    
    # Agriculture & Farming
    elif any(word in topic_lower for word in ['agriculture', 'farming', 'crop', 'livestock', 'agri']):
        refs = [
            f"Brady, N. C., & Weil, R. R. (2020). The nature and properties of soils (15th ed.). Pearson.",
            f"Reddy, S. R. (2019). Principles of agronomy (5th ed.). Kalyani Publishers.",
            f"Ensminger, M. E. (2021). Animal science (11th ed.). Interstate Publishers.",
            f"Rwanda TVET Board. (2023). {topic_clean} training standards. RTB."
        ]
    
    # Health & Nursing
    elif any(word in topic_lower for word in ['health', 'nursing', 'medical', 'patient', 'care']):
        refs = [
            f"Potter, P. A., & Perry, A. G. (2021). Fundamentals of nursing (10th ed.). Elsevier.",
            f"Kozier, B., & Erb, G. (2020). Fundamentals of nursing (11th ed.). Pearson.",
            f"Taylor, C., Lillis, C., & Lynn, P. (2019). Fundamentals of nursing (9th ed.). Wolters Kluwer.",
            f"Rwanda TVET Board. (2023). {topic_clean} competency framework. RTB."
        ]
    
    # Default for any other topic
    else:
        refs = [
            f"Rwanda Technical and Vocational Education and Training (TVET) Board. (2023). {topic_clean} training curriculum. RTB.",
            f"UNESCO-UNEVOC. (2021). TVET teaching and learning: A practical guide for {topic_clean}. UNESCO.",
            f"Ministry of Education Rwanda. (2022). {sector or 'Technical'} education standards for {topic_clean}. MINEDUC.",
            f"International Labour Organization. (2020). Skills development for {topic_clean}: A guide for TVET. ILO."
        ]
    
    return "\n".join(refs)

def generate_smart_objectives(topic, learning_outcomes):
    """Generate brilliant SMART objectives (Specific, Measurable, Achievable, Relevant, Time-bound)"""
    lo_lower = (learning_outcomes or '').lower()
    topic = topic or 'the topic'
    
    # Remove "Shield" from topic if present
    topic = topic.replace("Shield", "").replace("shield", "").strip()
    
    # RTB Format: Action Verb + Adverb + Content (as used in...)
    if any(word in lo_lower for word in ['create', 'design', 'develop', 'construct', 'produce', 'build']):
        objectives = [
            f"Define correctly key concepts and terminology of {topic}",
            f"Create accurately practical solutions applying {topic} principles",
            f"Produce effectively quality outputs demonstrating mastery of {topic}"
        ]
    elif any(word in lo_lower for word in ['analyze', 'evaluate', 'compare', 'examine', 'assess']):
        objectives = [
            f"Identify properly all components and elements of {topic}",
            f"Analyze correctly the underlying principles and theories of {topic}",
            f"Evaluate effectively different approaches and methods in {topic}"
        ]
    elif any(word in lo_lower for word in ['apply', 'demonstrate', 'implement', 'use', 'operate', 'perform']):
        objectives = [
            f"Define correctly {topic} terminology in professional practice",
            f"Demonstrate properly {topic} procedures following industry standards",
            f"Apply accurately {topic} techniques to solve real-world problems"
        ]
    else:
        objectives = [
            f"Define correctly fundamental concepts of {topic}",
            f"Identify properly key elements and components of {topic}",
            f"Explain accurately the principles and applications of {topic}"
        ]
    
    return "By the end of the session trainees will be able to\n" + "\n".join([f"✓ {obj}" for obj in objectives])

def generate_session_plan_docx(session_plan):
    """Generate RTB session plan - OFFICIAL FORMAT"""
    doc = Document()
    
    # LANDSCAPE orientation (official RTB)
    section = doc.sections[0]
    section.orientation = 1
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # Logo header (FIRST PAGE ONLY)
    rtb_logo = getattr(session_plan, 'rtb_logo_path', None)
    school_logo = getattr(session_plan, 'school_logo_path', None)
    
    if rtb_logo or school_logo:
        logo_table = doc.add_table(rows=1, cols=3)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # RTB Logo (left)
        if rtb_logo and os.path.exists(rtb_logo):
            try:
                logo_table.cell(0, 0).paragraphs[0].add_run().add_picture(rtb_logo, width=Inches(1.2))
                logo_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
        
        # School Logo (right)
        if school_logo and os.path.exists(school_logo):
            try:
                logo_table.cell(0, 2).paragraphs[0].add_run().add_picture(school_logo, width=Inches(1.2))
                logo_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
        
        doc.add_paragraph()  # Space after logos
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('SESSION PLAN')
    title_run.font.name = RTB_FONT
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    
    # Main table: 19 rows x 8 columns (OFFICIAL RTB - properly structured)
    table = doc.add_table(rows=19, cols=8)
    table.style = 'Table Grid'
    
    # Set line spacing to 1.5 for all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
    
    # Helper function
    def set_text(row, col, text, bold=True, colspan=1):
        cell = table.cell(row, col)
        if colspan > 1:
            cell.merge(table.cell(row, col + colspan - 1))
        para = cell.paragraphs[0]
        para.paragraph_format.line_spacing = 1.5  # Set 1.5 line spacing
        run = para.add_run(text)
        run.font.name = RTB_FONT
        run.font.size = Pt(12)
        run.font.bold = bold
        return cell
    
    # Row 0: Sector, Trade, Date
    set_text(0, 0, f"Sector: {session_plan.sector or ''}", colspan=2)
    set_text(0, 2, f"Trade: {session_plan.trade or ''}", colspan=4)
    set_text(0, 6, f"Date: {session_plan.date or ''}", colspan=2)
    
    # Row 1: Trainer, Term
    set_text(1, 0, f"Lead Trainer's name: {session_plan.trainer_name or ''}", colspan=5)
    set_text(1, 5, f"TERM: {session_plan.term or ''}", colspan=3)
    
    # Row 2: Module, Week, Trainees, Class
    set_text(2, 0, f"Module: {session_plan.module_code_title or ''}", colspan=3)
    set_text(2, 3, f"Week: {session_plan.week or ''}", colspan=2)
    set_text(2, 5, f"No. Trainees: {session_plan.number_of_trainees or ''}")
    set_text(2, 6, f"Class: {session_plan.class_name or ''}", colspan=2)
    
    # Row 3: Learning outcome
    set_text(3, 0, "Learning outcome:", colspan=2)
    set_text(3, 2, session_plan.learning_outcomes or '', bold=False, colspan=6)
    
    # Row 4: Indicative contents
    set_text(4, 0, "Indicative contents:", colspan=2)
    set_text(4, 2, session_plan.indicative_contents or '', bold=False, colspan=6)
    
    # Row 5: Topic
    set_text(5, 0, f"Topic: {session_plan.topic_of_session or ''}", colspan=8)
    
    # Row 6: Duration
    set_text(6, 0, f"Duration: {session_plan.duration or '40'} minutes", colspan=8)
    
    # Row 7: Objectives
    objectives_list = generate_smart_objectives(
        session_plan.topic_of_session or 'the topic',
        session_plan.learning_outcomes or ''
    )
    set_text(7, 0, "Objectives:", colspan=2)
    
    # Format objectives
    obj_cell = table.cell(7, 2)
    obj_cell.merge(table.cell(7, 7))
    
    # Split objectives and add
    obj_lines = objectives_list.split('\n')
    obj_para = obj_cell.paragraphs[0]
    
    # Add all lines with proper formatting
    for i, obj_line in enumerate(obj_lines):
        if i > 0:
            obj_para.add_run('\n')
        run = obj_para.add_run(obj_line)
        run.font.name = RTB_FONT
        run.font.size = Pt(12)
        run.font.bold = (i == 0)  # First line bold
    
    # Set line spacing for objectives
    obj_para.paragraph_format.line_spacing = 1.5
    
    # Row 8: Facilitation technique (with color)
    cell = set_text(8, 0, f"Facilitation technique: {session_plan.facilitation_techniques or 'Trainer Guided'}", colspan=8)
    set_cell_color(cell, "4472C4")  # Professional blue
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Row 9: Headers (with color - RTB standard gray)
    headers = [(9, 0, "Introduction"), (9, 1, "Activities"), (9, 6, "Resources"), (9, 7, "Duration")]
    for row, col, text in headers:
        if col == 1:
            cell = set_text(row, col, text, colspan=5)
        else:
            cell = set_text(row, col, text)
        set_cell_color(cell, "DBDBDB")  # RTB standard gray
    
    # Row 10: Introduction (with light blue background)
    topic = session_plan.topic_of_session or 'the topic'
    intro_activities = f"""Trainer's activity:
• Greets learners warmly and creates positive learning environment
• Gives trainees attendance list to sign
• Reviews previous session by asking 2-3 key questions
• Announces the new topic: \"{topic}\"
• Explains session objectives and expected outcomes
• Checks learners' prior knowledge about {topic}

Learner's activity:
• Replies to greetings enthusiastically
• Signs attendance list
• Responds to review questions from previous session
• Pays attention to new topic announcement
• Reads and understands session objectives
• Shares prior knowledge or experience with {topic}"""
    
    cell = set_text(10, 0, "Introduction")
    set_cell_color(cell, "D9E1F2")  # Light blue
    set_text(10, 1, intro_activities, bold=False, colspan=5)
    set_text(10, 6, "Attendance list\nPens\nWhiteboard\nMarkers", bold=False)
    set_text(10, 7, "10 min")
    
    # Row 11-14: Development (technique-specific with brilliant, detailed activities)
    technique = session_plan.facilitation_techniques or 'Trainer Guided'
    sector = session_plan.sector or 'General'
    
    # Determine resources based on sector
    if 'ICT' in sector.upper() or 'SOFTWARE' in (session_plan.trade or '').upper():
        dev_resources = "Computers/Laptops\nProjector/Smart board\nIDE Software\nInternet connection\nWhiteboard\nMarkers\nHandouts"
    elif 'HOSPITALITY' in sector.upper() or 'FOOD' in (session_plan.trade or '').upper():
        dev_resources = "Kitchen equipment\nIngredients\nProjector\nRecipe cards\nSafety equipment\nWhiteboard\nMarkers"
    elif 'CONSTRUCTION' in sector.upper() or 'BUILDING' in (session_plan.trade or '').upper():
        dev_resources = "Tools and equipment\nMaterials\nSafety gear\nBlueprints\nMeasuring instruments\nWhiteboard\nMarkers"
    else:
        dev_resources = "Relevant equipment\nProjector\nMaterials\nTools\nHandouts\nWhiteboard\nMarkers"
    
    if technique == 'Brainstorming':
        activities = f"""STEP 1: IDEA GENERATION PHASE (15 minutes)

Trainer's Activity:
• Presents the topic: \"{topic}\" with guiding question
• Creates safe environment for idea sharing
• Asks open-ended questions: \"What do you know about {topic}?\"
• Records ALL ideas on whiteboard without criticism or judgment
• Encourages wild ideas and creative thinking
• Ensures every learner contributes at least one idea
• Uses prompts: \"What else?\", \"Tell me more\"

Learner's Activity:
• Generates ideas freely about {topic}
• Shares thoughts without fear of judgment
• Builds on others' contributions (\"Yes, and...\")
• Participates actively in brainstorming
• Suggests creative and innovative ideas
• Listens to peers' contributions

STEP 2: IDEA ORGANIZATION AND CLUSTERING (15 minutes)

Trainer's Activity:
• Guides learners to group similar ideas together
• Helps identify patterns, themes, and categories
• Facilitates discussion on relationships between concepts
• Uses different colored markers for different categories
• Asks: \"Which ideas are related?\", \"What patterns do you see?\"
• Provides constructive feedback on idea clusters

Learner's activity:
• Categorizes ideas into logical groups
• Discusses connections between concepts
• Identifies key themes related to {topic}
• Organizes thoughts systematically
• Labels categories with appropriate names
• Explains reasoning for groupings

STEP 3: APPLICATION AND IMPLEMENTATION (30 minutes)

Trainer's Activity:
• Guides learners to select most relevant ideas
• Provides practical scenarios for applying {topic}
• Assigns groups to develop action plans
• Monitors application of brainstormed concepts
• Facilitates presentations of group solutions
• Provides feedback on practical applications

Learner's Activity:
• Applies organized ideas to solve real problems
• Works in groups on practical tasks using {topic} concepts
• Develops implementation strategies
• Demonstrates understanding through application
• Presents solutions to class
• Reflects on learning process"""
        resources = dev_resources + "\nSticky notes\nFlip charts\nColored markers"
        duration = "60 min"
    
    elif technique == 'Trainer Guided':
        activities = f"""STEP 1: DIRECT INSTRUCTION - \"I DO\" (15 minutes)

Trainer's Activity:
• Explains {topic} concepts, principles, and terminology clearly
• Demonstrates procedures and techniques step-by-step
• Presents real-world examples and applications of {topic}
• Shows complete process from start to finish
• Thinks aloud to model problem-solving approach
• Uses visual aids and demonstrations effectively
• Checks for understanding: \"Does everyone follow so far?\"

Learner's Activity:
• Follows and pays full attention to demonstration
• Takes detailed notes on key points and steps
• Observes demonstrations carefully
• Asks clarifying questions when needed
• Identifies critical steps in the process
• Makes connections to prior knowledge

STEP 2: GUIDED PRACTICE - \"WE DO\" (20 minutes)

Trainer's Activity:
• Works through examples WITH learners collaboratively
• Gives learners working examples on {topic}
• Asks learners to suggest next steps
• Guides practice activities with scaffolding
• Corrects errors using constructive feedback immediately
• Provides hints and prompts as needed
• Monitors learner progress continuously
• Gradually reduces support as confidence builds

Learner's activity:
• Works on provided examples with trainer guidance
• Practices {topic} procedures step-by-step
• Suggests solutions and next steps
• Displays work for feedback
• Compares their work with trainer's demonstration
• Makes corrections based on feedback
• Asks questions during practice
• Helps peers who are struggling

Step 3: Independent Practice - \"You Do\" (25 minutes)
Trainer's activity:
• Assigns independent practice tasks on {topic}
• Circulates to observe and support learners
• Provides individual assistance as needed
• Monitors for common errors or misconceptions
• Gives immediate, specific feedback
• Challenges advanced learners with extension tasks
• Ensures all learners achieve success

Learner's activity:
• Completes tasks independently
• Applies {topic} skills without direct guidance
• Tests and verifies their work
• Seeks help when genuinely stuck
• Demonstrates mastery of {topic}
• Helps peers after completing own work
• Reflects on learning progress"""
        resources = dev_resources
        duration = "60 min"
    
    elif technique == 'Group Discussion':
        activities = f"""Step 1: Topic Introduction and Group Formation (10 minutes)
Trainer's activity:
• Presents thought-provoking discussion question on {topic}
• Explains discussion objectives and expected outcomes
• Provides discussion guidelines and ground rules
• Divides learners into diverse small groups (4-6 members)
• Assigns roles: facilitator, recorder, timekeeper, presenter
• Provides discussion materials, resources, and guiding questions
• Sets clear time limits for each phase

Learner's activity:
• Listens carefully to the discussion question
• Understands discussion objectives and expectations
• Forms groups as directed
• Accepts assigned role and responsibilities
• Prepares materials for discussion
• Reviews guiding questions
• Begins thinking about {topic}

Step 2: Small Group Discussion (25 minutes)
Trainer's activity:
• Circulates among groups monitoring discussions
• Facilitates discussions without dominating
• Asks probing questions to deepen thinking
• Provides guidance and clarification when needed
• Ensures all members participate equally
• Redirects off-topic discussions
• Gives 5-minute warning before time ends
• Observes group dynamics and collaboration

Learner's activity:
• Discusses {topic} actively in small groups
• Shares ideas, experiences, and perspectives
• Listens to group members respectfully
• Records key points, conclusions, and insights
• Debates different viewpoints constructively
• Builds on others' ideas
• Stays focused on discussion question
• Ensures everyone contributes
• Prepares group presentation

Step 3: Whole Class Sharing and Synthesis (25 minutes)
Trainer's activity:
• Invites each group to present findings on {topic}
• Facilitates whole class discussion
• Asks clarifying and extension questions
• Highlights similarities and differences between groups
• Connects group findings to learning objectives
• Addresses misconceptions
• Summarizes key learning points
• Links discussion to real-world applications

Learner's activity:
• Presents group conclusions to class clearly
• Listens attentively to other groups' presentations
• Asks questions for clarification
• Provides constructive feedback
• Compares own group's ideas with others
• Participates in synthesizing class learning
• Takes notes on new insights
• Reflects on how discussion enhanced understanding of {topic}"""
        resources = dev_resources + "\nDiscussion guides\nFlip charts\nNotebooks"
        duration = "60 min"
    
    elif technique == 'Simulation':
        activities = f"""Step 1: Simulation Preparation and Setup (15 minutes)
Trainer's activity:
• Explains the simulation scenario for {topic} in detail
• Describes context, objectives, and expected outcomes
• Assigns roles with clear responsibilities and expectations
• Provides role cards, scenario sheets, and materials
• Explains rules, constraints, and success criteria
• Demonstrates expected behaviors and interactions
• Ensures learners understand their roles
• Sets up simulation environment
• Addresses questions and concerns

Learner's activity:
• Listens carefully to simulation scenario
• Understands the context and objectives
• Accepts and prepares for assigned role
• Reviews role card and responsibilities
• Studies materials and instructions
• Asks questions for clarification
• Mentally prepares for simulation
• Gathers necessary materials

Step 2: Simulation Execution (30 minutes)
Trainer's activity:
• Starts simulation with clear signal
• Monitors the simulation closely without interfering
• Observes learner performance and interactions
• Takes detailed notes on key observations
• Ensures safety and proper conduct
• Provides minimal guidance only if critical
• Manages time and keeps simulation on track
• Notes excellent examples and common errors
• Ends simulation at appropriate time

Learner's activity:
• Performs assigned role authentically and professionally
• Applies {topic} concepts in realistic context
• Interacts with peers according to scenario
• Makes decisions based on {topic} principles
• Stays in character throughout simulation
• Responds to challenges and problems
• Collaborates with team members
• Demonstrates skills and knowledge
• Maintains focus and engagement

Step 3: Debriefing and Reflection (25 minutes)
Trainer's activity:
• Leads structured reflection discussion
• Asks probing questions: \"What happened?\", \"Why?\", \"What if?\"
• Connects simulation experience to {topic} theory
• Highlights excellent performances
• Discusses challenges and how to overcome them
• Addresses misconceptions revealed during simulation
• Links to real-world professional practice
• Summarizes key learning points
• Reinforces correct application of {topic}

Learner's activity:
• Reflects deeply on simulation experience
• Shares observations, feelings, and insights
• Discusses challenges encountered and solutions tried
• Relates experience to {topic} concepts and theory
• Identifies what worked well and what didn't
• Learns from peers' experiences
• Identifies real-world applications
• Considers how to improve performance
• Connects simulation to career preparation"""
        resources = dev_resources + "\nSimulation materials\nRole cards\nScenario sheets\nProps"
        duration = "70 min"
    
    elif technique == 'Experiential Learning':
        activities = f"""Step 1: Concrete Experience - \"Doing\" (20 minutes)
Trainer's activity:
• Sets up engaging hands-on activity for {topic}
• Provides all necessary materials and tools
• Explains safety guidelines and procedures clearly
• Demonstrates proper techniques and procedures
• Facilitates the learning experience actively
• Monitors learner engagement and safety
• Encourages experimentation and exploration
• Allows learners to make mistakes and learn from them
• Provides minimal intervention

Learner's activity:
• Engages fully in hands-on activity
• Experiences {topic} directly through practice
• Observes and records observations carefully
• Follows safety procedures strictly
• Experiments with {topic} concepts
• Tries different approaches
• Makes mistakes and learns from them
• Collaborates with peers
• Stays curious and engaged

Step 2: Reflective Observation - \"Reflecting\" (20 minutes)
Trainer's activity:
• Guides learners to reflect on experience
• Asks probing questions: \"What did you notice?\", \"What surprised you?\"
• Encourages sharing of observations and discoveries
• Facilitates peer discussions and comparisons
• Helps identify patterns, insights, and lessons
• Creates safe space for honest reflection
• Records key insights on board
• Validates learners' experiences

Learner's activity:
• Reflects deeply on what happened during activity
• Discusses observations with peers
• Identifies what worked well and what didn't
• Analyzes experiences related to {topic}
• Shares feelings, reactions, and surprises
• Compares results with peers
• Identifies patterns and trends
• Asks \"why\" questions

Step 3: Abstract Conceptualization - \"Thinking\" (15 minutes)
Trainer's activity:
• Helps learners connect experience to theory
• Explains {topic} principles based on their experience
• Uses learners' examples to illustrate concepts
• Introduces relevant terminology and frameworks
• Helps form generalizations and rules
• Connects to professional standards

Learner's activity:
• Connects experience to {topic} theory
• Forms generalizations and conclusions
• Understands \"why\" things happened
• Learns proper terminology
• Develops conceptual understanding
• Takes notes on key principles

Step 4: Active Experimentation - \"Planning\" (15 minutes)
Trainer's activity:
• Provides opportunities to apply learning
• Assigns new challenges using {topic}
• Guides transfer to new situations
• Encourages creative application
• Monitors improved performance

Learner's activity:
• Applies {topic} concepts to new situations
• Practices skills in different contexts
• Tests new understanding
• Plans how to use learning in real work
• Demonstrates improved performance
• Shares strategies with peers"""
        resources = dev_resources + "\nHands-on materials\nTools\nSafety gear\nWorksheets"
        duration = "70 min"
    
    elif technique == 'Jigsaw':
        activities = f"""Step 1: Expert Group Formation (10 minutes)
Trainer's activity:
• Divides {topic} into 4-5 distinct subtopics
• Assigns learners to expert groups (one per subtopic)
• Provides comprehensive resources for each subtopic
• Explains expert group responsibilities and expectations
• Sets clear learning objectives for each subtopic
• Provides study guides and key questions
• Explains the jigsaw process clearly

Learner's activity:
• Joins assigned expert group
• Receives materials on specific subtopic
• Understands learning objectives
• Reviews resources provided
• Prepares to become expert on subtopic
• Discusses initial understanding with group

Step 2: Expert Group Study and Mastery (30 minutes)
Trainer's activity:
• Supports expert groups in studying subtopics
• Ensures deep understanding of content
• Monitors group progress and discussions
• Provides additional resources as needed
• Asks questions to check understanding
• Verifies mastery before jigsaw phase
• Helps groups prepare teaching strategies
• Ensures all members can teach the content

Learner's activity:
• Collaborates closely with expert group members
• Studies assigned subtopic of {topic} thoroughly
• Discusses and clarifies concepts together
• Prepares teaching materials and examples
• Masters content to teach others effectively
• Practices explaining key concepts
• Creates visual aids or notes
• Ensures everyone in group understands
• Prepares to answer questions

Step 3: Jigsaw Groups and Peer Teaching (30 minutes)
Trainer's activity:
• Forms jigsaw groups (one expert from each subtopic)
• Monitors teaching and learning process
• Circulates to ensure quality teaching
• Clarifies misconceptions when needed
• Ensures all members participate actively
• Manages time for each expert presentation
• Assesses overall understanding of {topic}
• Facilitates final synthesis

Learner's activity:
• Joins jigsaw group with diverse experts
• Teaches subtopic to jigsaw group members clearly
• Uses examples and explanations effectively
• Learns from other experts' presentations
• Asks questions for clarification
• Takes detailed notes on all subtopics
• Synthesizes complete understanding of {topic}
• Helps peers understand difficult concepts
• Completes full picture of {topic}

Step 4: Whole Class Synthesis (10 minutes)
Trainer's activity:
• Facilitates whole class discussion
• Connects all subtopics together
• Clarifies any remaining confusion
• Assesses complete understanding

Learner's activity:
• Participates in class synthesis
• Shares insights from jigsaw experience
• Asks final questions
• Demonstrates complete understanding of {topic}"""
        resources = dev_resources + "\nSubtopic materials\nHandouts\nNotebooks"
        duration = "80 min"
    
    else:
        activities = f"""Step 1: Content Introduction (15 minutes)
Trainer's activity:
• Presents {topic} concepts clearly and systematically
• Explains key principles and procedures
• Provides relevant examples and demonstrations
• Uses visual aids effectively
• Checks for understanding regularly

Learner's activity:
• Listens attentively
• Takes comprehensive notes
• Asks questions for understanding
• Engages with examples

Step 2: Guided Practice (25 minutes)
Trainer's activity:
• Guides practice activities on {topic}
• Provides support and feedback
• Monitors learner progress
• Adjusts pace as needed

Learner's activity:
• Practices {topic} skills with guidance
• Works through examples
• Applies concepts with support
• Seeks help when needed

Step 3: Independent Application (20 minutes)
Trainer's activity:
• Assigns independent tasks
• Observes and assists as needed
• Provides feedback

Learner's activity:
• Completes tasks independently
• Demonstrates understanding of {topic}
• Applies skills to solve problems"""
        resources = dev_resources
        duration = "60 min"
    
    cell = set_text(11, 0, "Development/Body")
    set_cell_color(cell, "D9E1F2")  # Light blue
    set_text(11, 1, activities, bold=False, colspan=5)
    set_text(11, 6, resources, bold=False)
    set_text(11, 7, duration)
    
    # Row 12: Conclusion (with light blue background)
    conclusion_text = f"""Trainer's activity:
• Summarizes key points of {topic} systematically
• Reviews main concepts covered during session
• Connects learning to session objectives
• Highlights most important takeaways
• Answers final questions and clarifications
• Links current session to next topic
• Previews upcoming learning

Learner's activity:
• Participates actively in summary
• Asks clarification questions
• Reflects on personal learning
• Makes connections to prior knowledge
• Identifies key concepts learned
• Prepares for next session"""
    cell = set_text(12, 0, "Conclusion")
    set_cell_color(cell, "D9E1F2")  # Light blue
    set_text(12, 1, conclusion_text, bold=False, colspan=5)
    set_text(12, 6, "Projector\nNotes\nWhiteboard", bold=False)
    set_text(12, 7, "10 min")
    
    # Row 13-15: Assessment (3 separate rows with light green background)
    summary_text = f"""Summary:
Together with trainees, trainer summarizes the session by asking targeted questions to check if objectives have been achieved. Questions focus on key concepts, procedures, and applications of {topic}. Trainees participate actively in the summary by providing immediate, thoughtful replies demonstrating their understanding."""
    cell = set_text(13, 0, summary_text, bold=False, colspan=8)
    set_cell_color(cell, "E2EFDA")  # Light green
    
    assessment_text = f"""Assessment/Assignment:
Trainer provides practical assessment tasks to evaluate understanding of {topic}. Assessment includes:
• Oral questions testing conceptual understanding
• Practical exercises applying {topic} skills
• Home-based assignment to master {topic} concepts independently
• Self-assessment checklist for learners
Trainees complete assessment tasks, demonstrating competency in {topic}."""
    cell = set_text(14, 0, assessment_text, bold=False, colspan=8)
    set_cell_color(cell, "E2EFDA")  # Light green
    
    evaluation_text = f"""Evaluation of the session:
With trainees, trainer evaluates the session by asking:
• \"What did you like most about today's session on {topic}?\"
• \"What helped you learn best?\"
• \"What should we improve for next sessions?\"
• \"How confident do you feel about {topic} now?\"

Trainees participate in evaluation by giving their honest views on what they liked and what to improve for future sessions. Trainer records feedback for continuous improvement.

Trainer links the current session to the next coming session, explaining how {topic} connects to upcoming topics. Trainees pay attention and prepare mentally for next learning."""
    cell = set_text(15, 0, evaluation_text, bold=False, colspan=8)
    set_cell_color(cell, "E2EFDA")  # Light green
    
    # Row 16: References (with light yellow background)
    refs = generate_apa_references(topic, session_plan.sector or '', session_plan.learning_outcomes or '')
    set_text(16, 0, "References:", colspan=2)
    cell = set_text(16, 2, refs, bold=False, colspan=6)
    set_cell_color(table.cell(16, 0), "FFF2CC")  # Light yellow
    set_cell_color(cell, "FFF2CC")  # Light yellow
    
    # Row 17: Appendices (with light yellow background)
    appendices_text = f"PowerPoint presentation on {topic}\nPractical worksheets and exercises\nAssessment rubrics and marking schemes\nAnswer keys and solutions\nAdditional reading materials\nVideo tutorials (if applicable)\nOnline resources and links"
    set_text(17, 0, "Appendices:", colspan=2)
    cell = set_text(17, 2, appendices_text, bold=False, colspan=6)
    set_cell_color(table.cell(17, 0), "FFF2CC")  # Light yellow
    set_cell_color(cell, "FFF2CC")  # Light yellow
    
    # Row 18: Reflection (with light yellow background)
    reflection_text = session_plan.reflection or f"Session on {topic} delivered successfully using {technique} approach. Learners showed excellent engagement and understanding. Key strengths: [To be filled after delivery]. Areas for improvement: [To be filled after delivery]. Overall, learning objectives were achieved effectively."
    set_text(18, 0, "Reflection:", colspan=2)
    cell = set_text(18, 2, reflection_text, bold=False, colspan=6)
    set_cell_color(table.cell(18, 0), "FFF2CC")  # Light yellow
    set_cell_color(cell, "FFF2CC")  # Light yellow
    
    # Save
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name

def generate_scheme_of_work_docx(scheme):
    """Generate brilliant RTB Scheme of Work - OFFICIAL FORMAT"""
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = 1
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # Logo header (RTB left, School right)
    rtb_logo = getattr(scheme, 'rtb_logo_path', None)
    school_logo = getattr(scheme, 'school_logo_path', None)
    
    if rtb_logo or school_logo:
        logo_table = doc.add_table(rows=1, cols=3)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # RTB Logo (left)
        if rtb_logo and os.path.exists(rtb_logo):
            try:
                logo_table.cell(0, 0).paragraphs[0].add_run().add_picture(rtb_logo, width=Inches(1.2))
                logo_table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            except:
                pass
        
        # School Logo (right)
        if school_logo and os.path.exists(school_logo):
            try:
                logo_table.cell(0, 2).paragraphs[0].add_run().add_picture(school_logo, width=Inches(1.2))
                logo_table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            except:
                pass
        
        doc.add_paragraph()  # Space after logos
    
    # School header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(6)
    run = header.add_run(f"{getattr(scheme, 'province', None) or 'PROVINCE'}\n")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = RTB_FONT
    run = header.add_run(f"{getattr(scheme, 'district', None) or 'DISTRICT'}\n")
    run.font.size = Pt(11)
    run.font.name = RTB_FONT
    run = header.add_run(f"{getattr(scheme, 'sector', None) or 'SECTOR'}\n")
    run.font.size = Pt(11)
    run.font.name = RTB_FONT
    run = header.add_run(f"{getattr(scheme, 'school', None) or 'SCHOOL'}")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = RTB_FONT
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run('SCHEME OF WORK')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = RTB_FONT
    
    # Module Details Table
    details_table = doc.add_table(rows=6, cols=4)
    details_table.style = 'Table Grid'
    
    def add_detail_row(row, col1_label, col1_value, col2_label, col2_value):
        cell = details_table.cell(row, 0)
        para = cell.paragraphs[0]
        run = para.add_run(col1_label)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = RTB_FONT
        
        cell = details_table.cell(row, 1)
        para = cell.paragraphs[0]
        run = para.add_run(col1_value)
        run.font.size = Pt(11)
        run.font.name = RTB_FONT
        
        cell = details_table.cell(row, 2)
        para = cell.paragraphs[0]
        run = para.add_run(col2_label)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = RTB_FONT
        
        cell = details_table.cell(row, 3)
        para = cell.paragraphs[0]
        run = para.add_run(col2_value)
        run.font.size = Pt(11)
        run.font.name = RTB_FONT
    
    add_detail_row(0, 'Sector:', getattr(scheme, 'sector', None) or '', 'Trainer:', scheme.trainer_name or '')
    add_detail_row(1, 'Trade:', getattr(scheme, 'department_trade', None) or '', 'School Year:', getattr(scheme, 'school_year', None) or '')
    add_detail_row(2, 'Qualification Title:', scheme.qualification_title or '', 'Term:', getattr(scheme, 'terms', None) or 'ALL TERMS')
    add_detail_row(3, 'RQF Level:', getattr(scheme, 'rqf_level', None) or '', 'Module Code & Title:', getattr(scheme, 'module_code_title', None) or '')
    add_detail_row(4, 'Date:', getattr(scheme, 'school_year', None) or '', 'Learning Hours:', getattr(scheme, 'module_hours', None) or getattr(scheme, 'module_learning_hours', None) or '')
    add_detail_row(5, 'Class Name:', getattr(scheme, 'class_name', None) or '', 'Number of Classes:', getattr(scheme, 'number_of_classes', None) or '')
    
    # Default values for standard RTB format
    DEFAULT_ACTIVITIES = '''• Demonstration and simulation
• Individual and group work
• Practical exercises
• Individualized learning
• Trainer guided instruction
• Group discussion
• Hands-on practice'''
    
    DEFAULT_RESOURCES = '''Computers/Equipment
Projector
Projection screen
Printers
Internet connection
Whiteboard and markers
Textbooks and handouts
Practical materials'''
    
    DEFAULT_ASSESSMENT = '''Written test
Practical assessment
Oral questions
Observation
Project work
Portfolio assessment
Peer assessment'''
    
    DEFAULT_PLACE = '''Classroom
Computer lab
Workshop
Practical area'''
    
    # Generate for each term
    for term_num in [1, 2, 3]:
        # Term heading
        term_para = doc.add_paragraph()
        term_para.paragraph_format.space_before = Pt(12)
        term_para.paragraph_format.space_after = Pt(6)
        term_run = term_para.add_run(f'TERM {term_num}')
        term_run.font.size = Pt(14)
        term_run.font.bold = True
        term_run.font.name = RTB_FONT
        
        # Get term data
        weeks = getattr(scheme, f'term{term_num}_weeks', None) or f'Week 1-12'
        los = getattr(scheme, f'term{term_num}_learning_outcomes', None) or ''
        ics = getattr(scheme, f'term{term_num}_indicative_contents', None) or ''
        duration = getattr(scheme, f'term{term_num}_duration', None) or '10 hours'
        
        # Create table with 9 columns (official RTB format)
        table = doc.add_table(rows=2, cols=9)
        table.style = 'Table Grid'
        
        # Row 0: Main headers with gray background
        headers = [
            (0, 0, 'Weeks', True),
            (0, 1, 'Competence code and name', False),
            (0, 4, 'Learning Activities', True),
            (0, 5, 'Resources (Equipment, tools, and materials)', True),
            (0, 6, 'Evidences of formative assessment', True),
            (0, 7, 'Learning Place', True),
            (0, 8, 'Observation', True)
        ]
        
        # Weeks (rowspan 2)
        cell = table.cell(0, 0)
        cell.merge(table.cell(1, 0))
        cell.text = 'Weeks'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_color(cell, 'D9D9D9')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = RTB_FONT
        
        # Competence code and name (cols 1-3 merged)
        cell = table.cell(0, 1)
        cell.merge(table.cell(0, 3))
        cell.text = 'Competence code and name'
        set_cell_color(cell, 'D9D9D9')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = RTB_FONT
        
        # Other headers (rowspan 2)
        for col, text in [(4, 'Learning Activities'), (5, 'Resources\n(Equipment, tools, and materials)'), 
                          (6, 'Evidences of formative assessment'), (7, 'Learning Place'), (8, 'Observation')]:
            cell = table.cell(0, col)
            cell.merge(table.cell(1, col))
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_color(cell, 'D9D9D9')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.name = RTB_FONT
        
        # Row 1: Sub-headers
        for col, text in [(1, 'Learning outcome (LO)'), (2, 'Duration'), (3, 'Indicative content (IC)')]:
            cell = table.cell(1, col)
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_color(cell, 'D9D9D9')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.name = RTB_FONT
        
        # Parse LOs and ICs
        lo_list = [l.strip() for l in los.split('\n') if l.strip()]
        ic_list = [c.strip() for c in ics.split('\n') if c.strip()]
        
        # Add data rows
        num_rows = max(len(lo_list), len(ic_list), 1)
        for i in range(num_rows):
            row = table.add_row()
            
            # Weeks (only in first row)
            if i == 0:
                cell = row.cells[0]
                para = cell.paragraphs[0]
                run = para.add_run(weeks)
                run.font.bold = True
                run.font.size = Pt(12)
                run.font.name = RTB_FONT
            
            # LO
            lo_text = lo_list[i] if i < len(lo_list) else ''
            cell = row.cells[1]
            para = cell.paragraphs[0]
            run_bold = para.add_run(f"LO{i+1}: ")
            run_bold.font.bold = True
            run_bold.font.size = Pt(12)
            run_bold.font.name = RTB_FONT
            run_normal = para.add_run(lo_text)
            run_normal.font.size = Pt(12)
            run_normal.font.name = RTB_FONT
            
            # Duration
            cell = row.cells[2]
            para = cell.paragraphs[0]
            run = para.add_run(duration)
            run.font.size = Pt(12)
            run.font.name = RTB_FONT
            
            # IC
            ic_text = ic_list[i] if i < len(ic_list) else ''
            cell = row.cells[3]
            para = cell.paragraphs[0]
            run_bold = para.add_run(f"IC{i+1}: ")
            run_bold.font.bold = True
            run_bold.font.size = Pt(12)
            run_bold.font.name = RTB_FONT
            run_normal = para.add_run(ic_text)
            run_normal.font.size = Pt(12)
            run_normal.font.name = RTB_FONT
            
            # Default values for last 5 columns
            for idx, text in [(4, DEFAULT_ACTIVITIES), (5, DEFAULT_RESOURCES), 
                              (6, DEFAULT_ASSESSMENT), (7, DEFAULT_PLACE), (8, '')]:
                cell = row.cells[idx]
                para = cell.paragraphs[0]
                run = para.add_run(text)
                run.font.size = Pt(12)
                run.font.name = RTB_FONT
        
        # Integrated Assessment row
        assess_row = table.add_row()
        cell = assess_row.cells[0]
        para = cell.paragraphs[0]
        run = para.add_run('Integrated Assessment (for specific module)')
        run.font.size = Pt(12)
        run.font.name = RTB_FONT
        run.font.bold = True
        
        for col, text in [(1, 'Practical task'), (2, 'Consumables and materials'), (7, 'Workshop/Lab')]:
            cell = assess_row.cells[col]
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(12)
            run.font.name = RTB_FONT
        
        # Trainer signature
        sig = doc.add_paragraph(f"Trainer's name and signature: {scheme.trainer_name or '_' * 50}")
        sig.paragraph_format.space_before = Pt(6)
        sig.paragraph_format.space_after = Pt(6)
        for run in sig.runs:
            run.font.size = Pt(12)
            run.font.name = RTB_FONT
        
        if term_num < 3:
            doc.add_page_break()
    
    # Final sign-offs
    doc.add_page_break()
    signoff = doc.add_paragraph()
    signoff.paragraph_format.space_before = Pt(12)
    run = signoff.add_run(f"Prepared by: (Name, position and Signature)\nTRAINER: {scheme.trainer_name or '_' * 50}\n\n")
    run.font.size = Pt(12)
    run.font.name = RTB_FONT
    run = signoff.add_run(f"Verified by: (Name, position and Signature)\nDOS: {getattr(scheme, 'dos_name', None) or '_' * 50}\n\n")
    run.font.size = Pt(12)
    run.font.name = RTB_FONT
    run = signoff.add_run(f"Approved by: (Name, position and Signature)\nSCHOOL MANAGER: {getattr(scheme, 'manager_name', None) or '_' * 50}")
    run.font.size = Pt(12)
    run.font.name = RTB_FONT
    
    # Save
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    return temp_file.name
