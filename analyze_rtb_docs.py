from docx import Document
import os

def analyze_rtb_document(file_path):
    """Analyze RTB official document structure"""
    print(f"\n{'='*60}")
    print(f"Analyzing: {os.path.basename(file_path)}")
    print('='*60)
    
    doc = Document(file_path)
    
    # Analyze sections
    print(f"\nSections: {len(doc.sections)}")
    for i, section in enumerate(doc.sections):
        print(f"  Section {i+1}:")
        print(f"    Orientation: {'Landscape' if section.orientation == 1 else 'Portrait'}")
        print(f"    Page width: {section.page_width.inches:.2f} inches")
        print(f"    Page height: {section.page_height.inches:.2f} inches")
        print(f"    Margins: L={section.left_margin.inches:.2f}, R={section.right_margin.inches:.2f}, T={section.top_margin.inches:.2f}, B={section.bottom_margin.inches:.2f}")
    
    # Analyze header
    if doc.sections[0].header:
        print(f"\nHeader content:")
        for para in doc.sections[0].header.paragraphs:
            if para.text.strip():
                print(f"  Text: {para.text[:100]}")
                if para.runs:
                    run = para.runs[0]
                    print(f"    Font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}")
    
    # Analyze tables
    print(f"\nTables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\n  Table {i+1}:")
        print(f"    Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        print(f"    Style: {table.style.name if table.style else 'None'}")
        
        # Sample first few cells
        if len(table.rows) > 0:
            print(f"    First row content:")
            for j, cell in enumerate(table.rows[0].cells[:3]):
                text = cell.text.strip()[:50]
                if text:
                    print(f"      Cell {j}: {text}")
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        run = cell.paragraphs[0].runs[0]
                        print(f"        Font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}")
    
    # Analyze paragraphs
    print(f"\nParagraphs: {len(doc.paragraphs)}")
    for i, para in enumerate(doc.paragraphs[:5]):
        if para.text.strip():
            print(f"  Para {i+1}: {para.text[:80]}")
            if para.runs:
                run = para.runs[0]
                print(f"    Font: {run.font.name}, Size: {run.font.size}, Bold: {run.font.bold}, Alignment: {para.alignment}")

# Analyze all RTB documents
docs_dir = "DOC TO REFER TO"
for filename in os.listdir(docs_dir):
    if filename.endswith('.docx') and not filename.startswith('~'):
        file_path = os.path.join(docs_dir, filename)
        try:
            analyze_rtb_document(file_path)
        except Exception as e:
            print(f"Error analyzing {filename}: {e}")

print("\n" + "="*60)
print("Analysis Complete!")
print("="*60)
