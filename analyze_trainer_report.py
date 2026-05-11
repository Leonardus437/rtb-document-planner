from docx import Document

doc = Document("DOC TO REFER TO/Trainer's Assessment Report Template .docx")

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"\nPara {i}: '{para.text}'")
        if para.runs:
            run = para.runs[0]
            print(f"  Font: {run.font.name}, Size: {run.font.size.pt if run.font.size else 'None'} pt, Bold: {run.font.bold}")

print("\n\n=== TABLE ===")
table = doc.tables[0]
print(f"Rows: {len(table.rows)}")
print(f"Columns: {len(table.rows[0].cells)}")

print("\n--- All Rows ---")
for row_idx, row in enumerate(table.rows):
    print(f"\nRow {row_idx}:")
    for col_idx, cell in enumerate(row.cells):
        text = cell.text.strip()
        if text:
            print(f"  Col {col_idx}: {text[:80]}")
